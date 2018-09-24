#! python3
# custom_word_invites.py
# usage: python custom_word_invites.py {invite_list} {styled_document_name.docx} {new_document_name.docx}

import docx
import sys

if len(sys.argv) > 2:
    invite_list = sys.argv[1]
    original_doc_name = sys.argv[2]
    new_doc_name = sys.argv[3]

    # obtains list of guests from user's text file
    with open(invite_list, 'r') as guest_list:
        guests = guest_list.read().split('\n')
        guests = guests[:-1]

    # doc = docx.opendocx(original_doc_name)
    doc = docx.Document()
    for name in guests:
        doc.add_paragraph('It would be a pleasure to have the company of', 'Title')
        doc.add_paragraph(name, 'Heading 5')
        doc.add_paragraph('at 11010 Memory Lane on the Evening of', 'Subtitle')
        doc.add_paragraph('April 1st', 'Heading 5')
        doc.add_paragraph('at 7 o\'clock', 'Subtitle')
        doc.add_page_break()

    doc.save(new_doc_name)
else:
    print('usage: python custom_word_invites.py {invite_list} {styled_document_name.docx} {new_document_name.docx}')
